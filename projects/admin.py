from django.contrib import admin
from django.http import HttpResponse
import shutil
import os
from docx import Document
from django.conf import settings
from .models import Project

# Export action function
def export_all_projects_action(modeladmin, request, queryset):
    # Define the master folder
    master_folder = os.path.join(settings.MEDIA_ROOT, 'all_projects')
    if os.path.exists(master_folder):
        shutil.rmtree(master_folder)  # Clear the folder if it exists
    os.makedirs(master_folder, exist_ok=True)

    # Export all selected projects
    for project in queryset:
        # Create a folder for each project
        project_folder = os.path.join(master_folder, f"{project.title}")
        os.makedirs(project_folder, exist_ok=True)

        # Save project details into a .docx file
        docx_path = os.path.join(project_folder, f"{project.title}_details.docx")
        document = Document()
        document.add_heading(project.title, level=1)
        document.add_paragraph(f"Description: {project.description}")
        document.add_paragraph(f"GitHub Link: {project.github_link}")
        document.add_paragraph(f"Department: {project.department}")
        document.add_paragraph(f"Faculty Members: {', '.join([str(faculty) for faculty in project.faculty_members.all()])}")
        document.add_paragraph(f"Team Members: {', '.join([str(member) for member in project.team_members.all()])}")
        document.save(docx_path)

         # Save the project image
        if project.image:  # Assuming your model has an ImageField named 'image'
            image_path = project.image.path
            if os.path.exists(image_path):
                shutil.copy(image_path, os.path.join(project_folder, os.path.basename(image_path)))

        # Copy associated files (if they exist)
        if project.document and os.path.exists(project.document.path):
            shutil.copy(project.document.path, os.path.join(project_folder, os.path.basename(project.document.name)))
        if project.video and os.path.exists(project.video.path):
            shutil.copy(project.video.path, os.path.join(project_folder, os.path.basename(project.video.name)))

    # Create a ZIP file for the master folder
    zip_path = f"{master_folder}.zip"
    shutil.make_archive(master_folder, 'zip', master_folder)

    # Serve the ZIP file as a download
    with open(zip_path, 'rb') as zip_file:
        response = HttpResponse(zip_file.read(), content_type='application/zip')
        response['Content-Disposition'] = 'attachment; filename=all_projects.zip'
        return response

# Admin panel configuration
@admin.register(Project)
class ProjectAdmin(admin.ModelAdmin):
    list_display = ['title', 'department', 'created_at']
    actions = [export_all_projects_action]  # Add the export action
